## -------------------------------------------------
## load required packages
## -------------------------------------------------
import streamlit as st
import pandas as pd
from utils.utils import save_dict_with_markdown_to_word
import logging
from streamlit_extras.app_logo import add_logo
from streamlit_extras.colored_header import colored_header
from langchain_community.vectorstores import AzureSearch
from langchain_community.embeddings import AzureOpenAIEmbeddings
from langchain.chains import RetrievalQA, ConversationalRetrievalChain
from langchain_community.chat_models import AzureChatOpenAI, ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter, SentenceTransformersTokenTextSplitter
import os
# from lib.common import get_encoding_name
from utils.process_file import process_file
from src.build_document_v0 import generate_response
from src.prepare_prompts import prepare_prompts
from utils.extract_sections import *
from docxtpl import DocxTemplate, RichText
from datetime import datetime
import jinja2
import io
import spacy
# nlp = spacy.load("en_core_web_lg")
from docx import Document
from datetime import date
import time
import string
import random
import ast
import json
import re
import pypandoc
import base64

## -------------------------------------------------
## Configure logging
## -------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

## -------------------------------------------------
## env vars and required vars
## -------------------------------------------------
os.environ['LANGCHAIN_TRACING_V2'] = 'true'
os.environ['LANGCHAIN_ENDPOINT'] = os.getenv("LANGCHAIN_ENDPOINT")
os.environ['LANGCHAIN_API_KEY'] = os.getenv("LANGCHAIN_API_KEY")
os.environ['LANGCHAIN_PROJECT'] = os.getenv("LANGCHAIN_PROJECT")

vector_store_address: str = os.getenv("SEARCH_ENDPOINT")
vector_store_password: str = os.getenv("SEARCH_KEY")

azure_openai_api_key: str = os.getenv("OPENAI_API_KEY_AZURE")
azure_endpoint: str = os.getenv("OPENAI_ENDPOINT_AZURE")

openai_api_key: str = os.getenv("OPENAI_API_KEY")

logger.info("Environment variables loaded")

## -------------------------------------------------
## set page config and options
## -------------------------------------------------
st.cache_data.clear()
st.cache_resource.clear()
logger.info("Cleared app cache")

st.set_page_config(layout="wide", page_title="Generate RFP Response", page_icon="👷")
add_logo(r"./synoptek-new-removebg-3.png")
colored_header(
        label="👷 Generate RFP Response", description="\n",
        color_name="violet-70",
    )
logger.info("Page configuration and options set up")

## -------------------------------------------------
## session state initialization
## -------------------------------------------------
if 'clientOrg' not in st.session_state:
    st.session_state['clientOrg'] = ''

if 'docTitle' not in st.session_state:
    st.session_state['docTitle'] = ''

if 'vector_store' not in st.session_state:
    st.session_state['vector_store'] = None

if 'template' not in st.session_state:
    st.session_state['template'] = ''

if 'sections' not in st.session_state:
    st.session_state['sections'] = None

if 'section_dict' not in st.session_state:
    st.session_state['section_dict'] = None

if 'prepared_prompts' not in st.session_state:
    st.session_state['prepared_prompts'] = None

if 'generated_resp' not in st.session_state:
    st.session_state['generated_resp'] = None

if 'doc_generated' not in st.session_state:
    st.session_state['doc_generated'] = False

if 'bio' not in st.session_state:
    st.session_state['bio'] = None

if 'case_studies' not in st.session_state:
    st.session_state['case_studies'] = None

logger.info("Initialized session state")

## -------------------------------------------------
## azure openai deployment and setup
## -------------------------------------------------
@st.cache_resource
def azure_openai_setup(azure_openai_api_key, azure_endpoint):
    """
    Set up Azure OpenAI models for QA and response generation.
    
    Args:
        azure_openai_api_key (str): API key for Azure OpenAI.
        azure_endpoint (str): Azure OpenAI endpoint.

    Returns:
        tuple: AzureChatOpenAI instances for QA and response generation.
    """
    try:
        logger.info("Setting up Azure OpenAI")
        deployment_gpt4_turbo = 'gpt-4-1106-preview'
        deployment_gpt4 = 'gpt-4'
        deployment_gpt35 = 'gpt-3.5-turbo-1106'
        deployemnt_gpt4_azure = 'gpt-4-aims'
        deployemnt_gpt4o_azure = 'gpt-4o'
        deployemnt_gpt35_azure = 'gpt-35-aims'

        llm_azure_qa = AzureChatOpenAI( 
            model_name=deployemnt_gpt4o_azure,
            openai_api_key=azure_openai_api_key,
            azure_endpoint=azure_endpoint,
            openai_api_version="2024-04-01-preview",
            temperature=0,
            max_tokens=4000,
            model_kwargs={'seed':123}
            ) 

        llm_azure_resp = AzureChatOpenAI( 
            model_name=deployemnt_gpt4o_azure,
            openai_api_key=azure_openai_api_key,
            azure_endpoint=azure_endpoint,
            openai_api_version="2024-04-01-preview",
            temperature=0,
            max_tokens=4000,
            model_kwargs={'seed':123}
            ) 

        logger.info("Azure OpenAI setup completed")
        return llm_azure_qa, llm_azure_resp
    except Exception as e:
        logger.exception("Error setting up Azure OpenAI: %s", e)
        st.error("Failed to set up Azure OpenAI. Please check the logs for details.")

azure_qa, azure_resp = azure_openai_setup(azure_openai_api_key, azure_endpoint)

# @st.cache_resource
# def openai_setup(openai_api_key):
#     """
#     Set up OpenAI models for QA and response generation.
    
#     Args:
#         openai_api_key (str): API key for OpenAI.

#     Returns:
#         tuple: ChatOpenAI instances for QA and response generation.
#     """
#     try:
#         logger.info("Setting up OpenAI")
#         deployment_gpt4_turbo = 'gpt-4-1106-preview'
#         deplyment_gpt4o = 'gpt-4o'
#         deployment_gpt4 = 'gpt-4'
#         deployment_gpt35 = 'gpt-3.5-turbo-1106'
#         deployemnt_gpt4_azure = 'gpt-4-aims'
#         deployemnt_gpt35_azure = 'gpt-35-aims'

#         llm_openai = ChatOpenAI(
#         model_name=deployment_gpt4,
#         openai_api_key=openai_api_key,
#         temperature=0,
#         max_tokens=4000,
#         model_kwargs={'seed':123}
#         )

#         logger.info("OpenAI setup completed")
#         return llm_openai
    
#     except Exception as e:
#         logger.exception("Error setting up OpenAI: %s", e)
#         st.error("Failed to set up OpenAI. Please check the logs for details.")
    
# llm_openai_ = openai_setup(openai_api_key)

## -------------------------------------------------
## azure cognitive search vector index setup
## -------------------------------------------------
@st.cache_resource
def azureai_search_setup(azure_endpoint, azure_openai_api_key):
    """
    Set up Azure AI search with embeddings.

    Args:
        azure_endpoint (str): Azure search endpoint.
        azure_openai_api_key (str): API key for Azure OpenAI.

    Returns:
        AzureOpenAIEmbeddings: Embeddings instance.
    """
    try:
        logger.info("Setting up Azure AI search")
        azure_deployment='embeddings-aims'
        embeddings = AzureOpenAIEmbeddings(
            azure_deployment=azure_deployment,
            openai_api_version="2024-04-01-preview",
            azure_endpoint=azure_endpoint,
            api_key=azure_openai_api_key,
        )
        logger.info("Azure AI search setup completed")
        return embeddings
    
    except Exception as e:
        logger.exception("Error setting up Azure AI search: %s", e)
        st.error("Failed to set up Azure AI search. Please check the logs for details.")

embedds = azureai_search_setup(azure_endpoint, azure_openai_api_key)

## -------------------------------------------------
## function for processing and uploading RFP request to Azure vector index
## -------------------------------------------------
@st.cache_resource
def create_vector_index(clientOrg, rfp, vector_store_address, vector_store_password):
    """
    Create a vector index for the given client organization and upload the RFP document.

    Args:
        clientOrg (str): Client organization name.
        rfp (UploadedFile): RFP document.
        vector_store_address (str): Vector store address.
        vector_store_password (str): Vector store password.

    Returns:
        AzureSearch: Vector store instance.
    """
    logger.info("Creating vector index for client organization: %s", clientOrg)
    index_name = "rfp-request-" + str(clientOrg).lower()
    try:
        vector_store_ = AzureSearch(
            azure_search_endpoint=vector_store_address,
            azure_search_key=vector_store_password,
            index_name=index_name,
            embedding_function=embedds.embed_query,
        )
        
        doc_count = vector_store_.client.get_document_count()
        logger.debug("Document count in vector store: %d", doc_count)

        if doc_count == 0:
            st.info("No document found in database for the given client.")
            if rfp is not None:
                logger.debug("Document provided, uploading to database.")
                st.info("Document attached. Uploading in database.")
                start_time = time.time()
                pages = process_file(file_name=rfp.name, file=rfp, bytes_data=rfp.read())
                # # from_tiktoken_encoder enables use to split on tokens rather than characters
                # recursive_text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                # chunk_size=300, 
                # chunk_overlap=125
                # )

                # recursive_text_splitter_chunks = recursive_text_splitter.split_documents(pages)
                
                vector_store_.add_documents(documents=pages)
                time_taken = round((time.time() - start_time) / 60, 2)
                st.success(f"Document uploaded in {time_taken} mins.")
                logger.info("Document uploaded to vector store %s", index_name)
                #check if the file exists with path.exists()
                if os.path.exists(r'./temp/'+rfp.name):
                    os.remove(r'./temp/'+rfp.name)
                    logger.info('file deleted from temp folder')
                else:
                    logger.info("File does not exists in temp folder")
            else:
                st.warning("No document provided to upload.")
                logger.warning("No document provided for uploading to vector store %s", index_name)
        else:
            st.info("Document already exists in the database.")
            logger.warning("Document already exists in the vector store %s", index_name)

        return vector_store_
    except Exception as e:
        error_message = f"An error occurred while creating vector index for {clientOrg}: {e}"
        logger.exception(error_message)
        st.error(error_message)

## -------------------------------------------------
## function to convert dict to markdown
## -------------------------------------------------
# def dict_to_markdown(d):
#     markdown_text = ""
#     for key, value in d.items():
#         markdown_text += f"### {key}\n\n{value}\n\n"
#     return markdown_text

## -------------------------------------------------
## upload rfp request
## -------------------------------------------------
with st.sidebar:
    
    st.session_state['clientOrg'] = st.text_input("**What is the client name?** 🚩")

    st.session_state['docTitle'] = st.text_input("**Provide title for the response document** 🚩")

    rfp = st.file_uploader("**Upload RFP request to database** 🚩")

    if st.session_state['clientOrg']:
        st.session_state['vector_store'] = create_vector_index(st.session_state['clientOrg'], rfp, vector_store_address, vector_store_password)
    else:
        st.warning("Add client name above")

## -------------------------------------------------
## Generate RFP App
## -------------------------------------------------
col1, col2 = st.columns([2, 3.5])
with col1:
    with st.container(border=True):
        # Define a function to handle the button click
        try:
            if st.button("Get Response Template", key='button1'):
                # st.write(st.session_state)
                # logger.info("Session State: %s", st.session_state)
                vector_store = st.session_state['vector_store']
                if vector_store:
                    ## set up QA retrieval
                    qa_chain = RetrievalQA.from_chain_type(
                        llm=azure_qa,
                        retriever=vector_store.as_retriever(search_kwargs={"k": 15,
                                                                           }),
                        return_source_documents=True
                    )

                if st.session_state['clientOrg']:
                    
                    st.info("Response template given in RFP request:")
                    template = qa_chain({"query": """
                                        -Extract document related to Technical Proposal from the contents of proposals and submittals.
                                        -Extract the submission guidelines from this RFP, including the format, required documents, and any specific instructions.
                                        -If above document not found then check for vendor response specifications.
                                        -If above document not found then check for response submission format or response requirements.
                                        -If above document not found then check for response requirements.
                                        -If above document not found then check for question and answers which outlines response template.
                                        -There are chances you can miss information as it might be on different pages.
                                        -So validate if information is complete, if not redo above steps again till you get all information.
                                        -Provide a detailed breakdown of each section required in the RFP response, including sub-sections.
                                        -List all the documents and materials that need to be submitted as per this RFP."""})
                    cleaned_text = re.sub(r'\[doc \d\.\d-\d\]', '', template['result'].lower())
                    st.session_state['template'] = cleaned_text    
                    logger.info("Response template generated for client organization: %s", st.session_state['clientOrg'])
                else:
                    logger.warning("Client organization name is missing when attempting to get the response template")
            
        except:
            st.error("Connection aborted. Please generate response template again")
            logger.exception("message")
        else:
            st.write(st.session_state['template'])

    
with col2:
    with st.container(border=True):
        st.info("Edit, select for deletion, and click + to add new sections here.")
        try:
            if st.session_state['template']:
                extract_sections_prompt = extract_sections(response_template = st.session_state['template'])
                extracted_sections = azure_qa.invoke(extract_sections_prompt)
                extract_sections_ = extracted_sections.content
                # extract_sections_ = extract_sections_.strip(" ")
                extract_sections_ = extract_sections_.strip("```")
                match = re.search(r'\[(.*?)\]', extract_sections_, re.DOTALL)
                if match:
                    list_str = match.group(0)
                    extracted_list = ast.literal_eval(list_str)
                    if 'Context' in extracted_list:
                        value_index =  extracted_list.index('Context')
                        extracted_list.insert(value_index, 'title page')
                        extracted_list = [ele.upper() for ele in extracted_list]
                        st.session_state['sections'] = extracted_list
                        # st.info("Selected Sections")
                        logger.info("Selected Sections")

                    else:
                        extracted_list = [ele.upper() for ele in extracted_list]
                        st.session_state['sections'] = extracted_list
                        # st.info("Selected Sections")
                        logger.info("Selected Sections")
                        

            if st.session_state['sections']:
                initial_sections = st.session_state['sections'][:]  # Make a copy of the original sections
                # print(initial_sections)
                section_df = pd.DataFrame({'Sections': initial_sections})
                section_df['User Inputs'] = None
                # print(section_df)
                edited_df = st.data_editor(data=section_df, 
                                                num_rows="dynamic", use_container_width=True)
                edited_df_list = edited_df.to_dict(orient="records")
                # print("edited_df_list",edited_df_list)

                edited_sections = [item['Sections'].strip() for item in edited_df_list if item['Sections'].strip()]
                user_inputs = [item['User Inputs'].strip() if item['User Inputs'] is not None else item['User Inputs'] for item in edited_df_list]

                section_dict = {section: input_ for section, input_ in zip(edited_sections, user_inputs)}

                new_sections = [section for section in edited_sections if section not in initial_sections]
                deleted_sections = [section for section in initial_sections if section not in edited_sections]

                # Update the session state based on changes
                if new_sections or deleted_sections:
                    if deleted_sections:
                        st.success(f"Deleted sections: {', '.join(deleted_sections)}")
                        logger.info(f"Deleted sections: {', '.join(deleted_sections)}")
                    if new_sections:
                        st.success(f"New sections added: {', '.join(new_sections)}")
                        logger.info(f"New sections added: {', '.join(new_sections)}")
                        # Render a text area for each new section
                        # for section in new_sections:
                        #     additional_info = st.text_area(f"Input for {section}", key=f'text_area_{section}')
                        #     section_dict[section] = additional_info  # Update the dictionary with the input
                
                    if section_dict:
                        on = st.toggle("Synoptek Sections")
                        if on:
                            default_synoptek_sections = ['SYNOPTEK OVERVIEW','SYNOPTEK\'s CULTURE AND APPROACH TO TALENT MANAGEMENT',
                                        'CASE STUDIES','QUALITY SECURITY AND COMPLIANCE',
                                        'ASSUMPTIONS AND CLIENT RESPONSIBILITIES']
                            
                            sections = st.multiselect(label = '## **What sections to be included?**',
                                       options = ['SYNOPTEK OVERVIEW','SYNOPTEK\'s CULTURE AND APPROACH TO TALENT MANAGEMENT', 'CASE STUDIES',
                                        'QUALITY SECURITY AND COMPLIANCE',
                                        'ASSUMPTIONS AND CLIENT RESPONSIBILITIES'],
                                        default=default_synoptek_sections, 
                                        key='sections0')
                            
                            all_select = st.checkbox('Select All Case Studies', key='all_select')
                            # case_studies = st.multiselect("## **Select case studies from here**",
                            case_studies= ['Health Care', 'Finance', 'Business Application', 'Business Intelligence', 'Production Development', 'Microsoft Azure', 'Workforce Productivity', 'Security', 'Marketing Media Company', 'Window Manufacturer Sales Forecast', 'Industrial Motor Manufacturer Predictive Maintenance', 'Logistics Carrier Organization', 'Recall', 'Shop Goodwill', 'Fastmore', 'National Insurance Company', 'Real Estate Transaction Company', 'Patientcare Solutions Provider', 'Marketing Technology Firm', 'Research and Advisory Firm', 'Non-Profit Organization', 'National Engineering Company', 'Toy Retailer', 'Sporting Goods Retailer', 'Modular Kitchen Manufacturer and Retailer', 'Beauty Services and Products Provider', 'Digital Media Company', 'Consumer Media Company', 'Luxury Retailer', 'Global Industrial Packaging Leader', 'Global Air Conditioning Manufacturer', 'SaaS Product Marketing', 'Research Advisory Firm', 'Transport Management System', 'Entertainment Company', 'Legacy App Modernization', 'Non-Profit CRM Integration', 'Healthcare BI Solution', 'Claims Processing App', 'E-commerce Strategy', 'Program Management']#,key='caseStudies0')
                            
                            if all_select:
                                cases = st.multiselect('## **What case studies to be included?**', case_studies, case_studies, key='caseStudies0')
                            else:
                                cases = st.multiselect('## **What case studies to be included?**', case_studies, key='caseStudies0')

                            st.session_state['case_studies'] = cases
                            
                            synoptek_section_dict = {sec: "" for sec in sections}
                            section_dict.update(synoptek_section_dict)
                            print("section_dict", section_dict)

                else:
                    st.success("No changes made to the sections")
                    logger.info("No changes made to the sections")

                
                
                if st.button('Build RFP Response', key='button2'):
                    st.session_state['section_dict'] = section_dict
                    # st.write(st.session_state['section_dict'])

            if st.session_state['section_dict']:
                # if st.button("Prepare prompts and generate RFP response"):
                prepared_prompts = prepare_prompts(sections=st.session_state['section_dict'],
                            resp_template=st.session_state['template'],
                            llm_resp=azure_qa)
                st.session_state['prepared_prompts'] = prepared_prompts
                st.success("Prompts Prepared")
                logger.info("Prompts Prepared")
                # st.write(st.session_state['prepared_prompts'])
                
                if st.session_state['prepared_prompts']:
                    
                    generated_resp = generate_response(vector_store=st.session_state['vector_store'],
                                                        llm_qa=azure_qa, llm_resp=azure_resp,
                                                        prompts=prepared_prompts, 
                                                        clientOrg=st.session_state['clientOrg'],
                                                        case_studies=st.session_state['case_studies'])
                    # st.write('\n\n'.join([key+'\n', values+'\n']) for key, values in generated_resp.items())
                    st.session_state['generated_resp'] = generated_resp
                    st.success("Generated response")
                    logger.info("Generated response")
        except:
            st.error("Connection aborted. Please generate response template again")
            logger.exception("message")
        # st.write(st.session_state)
        logger.info("Session State: %s", st.session_state)


if st.session_state['generated_resp']:
    try:
        st.write(st.session_state['generated_resp'])
        ref_doc = r'./reference_docs/reference_doc_v1.docx'

        rfp_resp = st.session_state['generated_resp']
        print("rfp_resp",type(rfp_resp))
        print("rfp_resp_values",type(rfp_resp.values()))

        cleaned_rfp_resp = {key:re.sub(r'\[doc \d\.\d-\d\]', '', str(values)) for key, values in rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?Conclusion,\s', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?Conclusion\b.*(?:\n.*)*', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?conclusion,\s', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?conclusion\b.*(?:\n.*)*', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'^Conclusion\b.*(?:\n.*)*', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        # ---
        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?Summary,\s', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?Summary\b.*(?:\n.*)*', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?summary,\s', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'\b(In\s)?summary\b.*(?:\n.*)*', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        cleaned_rfp_resp = {key:re.sub(r'^Summary\b.*(?:\n.*)*', '', str(values)) for key, values in cleaned_rfp_resp.items()}

        # cleaned_text = re.sub(r'\[doc \d\.\d-\d\]', '', template['result'].lower())

        save_dict_with_markdown_to_word(dictionary = cleaned_rfp_resp, 
                                        file_path = r'./output_files/'+st.session_state['clientOrg']+'.docx',
                                        title=st.session_state['docTitle'].upper())

        # code for enable document download
        doc = Document(r'./output_files/'+st.session_state['clientOrg']+'.docx')
        bio = io.BytesIO()
        doc.save(bio)
    
        st.session_state['bio'] = bio.getvalue()
        st.session_state['doc_generated'] = True

        # response = '\n\n'.join([values for key, values in st.session_state['generated_resp'].items()])
        
        # tpl = DocxTemplate(ref_doc)
        # context = {
        #             "date_today": datetime.now().strftime(format='%B %d, %Y'),
        #             "mspOrg":"SYNOPTEK LLC",
        #             "clientOrg":st.session_state["clientOrg"],
        #             "title": st.session_state["docTitle"],
        #             'response': response,
        #             }
        # jinja_env = jinja2.Environment(autoescape=True)
        # tpl.render(context, jinja_env)
        # tpl.save(r'./output_files/'+st.session_state['clientOrg']+'.docx')

        # doc = Document()

        # # Add a title to the document
        # doc.add_heading('Content', level=1)

        # # Iterate over the dictionary items and add them to the document
        # for key, value in st.session_state['generated_resp'].items():
        #     doc.add_paragraph(f'{key}: {value}')

        # doc.save(r'./output_files/'+st.session_state['clientOrg']+'.docx')
        
    except:
        st.sidebar.error("Connection aborted. Please generate response template again")
        logger.exception("message")

if st.session_state['doc_generated']:
    # st.sidebar.download_button(
    #     label="Click here to download",
    #     data=st.session_state['bio'],
    #     file_name=st.session_state['clientOrg']+'.docx',
    #     mime="docx"
    # ) # "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    # Convert the BytesIO to a base64 string
    b64 = base64.b64encode(st.session_state['bio']).decode()

    # Create a download link
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{st.session_state["clientOrg"]}.docx">Click here to download</a>'
    st.sidebar.markdown(href, unsafe_allow_html=True)

    # After the download button is clicked, delete the file
    if os.path.exists(r'./output_files/'+st.session_state['clientOrg']+'.docx'):
        os.remove(r'./output_files/'+st.session_state['clientOrg']+'.docx')

logger.info("-------------------------------------------------------------------")
