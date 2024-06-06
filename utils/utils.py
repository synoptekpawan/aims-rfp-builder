from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup

def html_to_docx(html, doc):
    """
    Helper function to convert HTML content to a Word document.
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    for element in soup:
        if element.name == 'p':
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em':
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = p.add_run(child.get_text())
                    run.font.name = 'Courier New'
                elif child.name:
                    run = p.add_run(child.get_text())
                else:
                    run = p.add_run(str(child))
        elif element.name == 'table':
            table = doc.add_table(rows=0, cols=0)
            for row in element.find_all('tr'):
                cells = row.find_all(['td', 'th'])
                if table.rows:
                    row_cells = table.add_row().cells
                else:
                    row_cells = table.add_row().cells
                    table.add_row()._element.getparent().remove(table.rows[0]._element) # Remove extra row

                for i, cell in enumerate(cells):
                    if len(row_cells) <= i:
                        table.add_column()
                    if row_cells[i].text:
                        row_cells = table.add_row().cells
                    row_cells[i].text = cell.get_text(strip=True)
                    if cell.name == 'th':
                        row_cells[i].paragraphs[0].runs[0].bold = True
        elif element.name:
            doc.add_paragraph(element.get_text())


def save_dict_with_markdown_to_word(dictionary, file_path, title):
    # Create a new Document
    doc = Document() # r'../reference_docs/reference_doc_v1.docx'
    
    # Add a title to the document
    doc.add_heading(title, level=1)
    
    # Iterate over the dictionary items and add them to the document
    for key, value in dictionary.items():
        doc.add_heading(key, level=2)
        html_content = markdown.markdown(value)
        html_to_docx(html_content, doc)
    
    # Save the document
    doc.save(file_path)